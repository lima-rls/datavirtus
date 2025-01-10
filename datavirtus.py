import os.path

import pandas as pd
import docx
import faker
import re
import json
from docx.shared import Inches


def cpf_or_cnpj(cpf_cnpj):
    cpf_cnpj = ''.join(filter(str.isdigit, str(cpf_cnpj)))
    if len(cpf_cnpj) == 11:
        return 'CPF'
    elif len(cpf_cnpj) == 14:
        return 'CNPJ'
    else:
        return 'Inválido'


def pseudoanonimizar(dataframe, coluna, tipo, seed=None):

    if seed is not None:
        faker.Faker.seed(seed)

    fake = faker.Faker('pt_BR')

    reais_unicos = dataframe[coluna].unique()
    if tipo == 'nome':
        fake_values = list(map(lambda _: fake.name(), range(len(reais_unicos))))
    elif tipo == 'email':
        fake_values = list(map(lambda _: fake.email(), range(len(reais_unicos))))
    elif tipo == 'cpf_cnpj':
        fake_values = []
        for valor in reais_unicos:
            tipo_documento = cpf_or_cnpj(valor)
            if tipo_documento == 'CPF':
                fake_values.append(fake.cpf())
            elif tipo_documento == 'CNPJ':
                fake_values.append(fake.cnpj())
    elif tipo == 'telefone':
        fake_values = list(map(lambda _: fake.phone_number(), range(len(reais_unicos))))
    elif tipo == 'endereco':
        fake_values = list(map(lambda _: fake.address(), range(len(reais_unicos))))
    elif tipo == 'data':
        fake_values = list(map(lambda _: fake.date(), range(len(reais_unicos))))
    else:
        raise ValueError('Tipo inválido')

    fake_dict = dict(zip(reais_unicos, fake_values))
    dataframe[coluna] = dataframe[coluna].replace(fake_dict)

    return dataframe


class RelatorioVirtus:

    def __init__(self, modelo_relatorio, arquivo_tags=None):
        self.modelo_relatorio = modelo_relatorio
        self.nome_modelo = modelo_relatorio.split('.')[0]
        self.doc = docx.Document(modelo_relatorio)

        if arquivo_tags is not None:
            self.mapa_tags = self.carregar_tags(arquivo_tags)
        else:
            self.mapa_tags = None
            self.arquivo_tags = self.nome_modelo + '_tags.json'



    def carregar_doc(self):
        try:
            return docx.Document(self.modelo_relatorio)
        except FileNotFoundError:
            raise FileNotFoundError("Arquivo de modelo de relatório não encontrado.")
        except Exception as e:
            raise Exception(f"Erro ao carregar o arquivo de modelo de relatório: {e}")


    def extrai_texto(self):
        texto = []
        for paragrafo in self.doc.paragraphs:
            texto.append(paragrafo.text)
        return texto

    def extrair_tags(self):
        tags = set()
        texto = self.extrai_texto()
        for paragraph in texto:
            tags.update(re.findall(r'(\|.*?\|)', paragraph))
        mapa_tags = {tag: None for tag in tags}
        return mapa_tags

    @staticmethod
    def ordenar_dict(dicionario):
        return dict(sorted(dicionario.items(), key=lambda x: x[0]))

    def exportar_tags(self, arquivo_saida=None):
        if arquivo_saida is None:
            arquivo_saida = self.nome_modelo + '_tags.json'

        if os.path.exists(arquivo_saida):
            raise FileExistsError('Arquivo de tags já existe. Carregue o arquivo existente ou altere o nome desejado.')

        mapa_tags = self.extrair_tags()

        with open(arquivo_saida, 'w', encoding='utf-8') as f:
            json.dump(self.ordenar_dict(mapa_tags), f, ensure_ascii=False, indent=4)
            print(f'Arquivo de tags {arquivo_saida} exportado com sucesso.')
            print(f'Use |IMAGEM...| para inserir um gráfico e |TABELA...| para inserir uma tabela. Demais tags |...| são para texto.')


    def carregar_tags(self, arquivo_tags):
        try:
            with open(arquivo_tags, 'r', encoding='utf-8') as f:
                self.mapa_tags = json.load(f)
                self.arquivo_tags = arquivo_tags
                print(f'Arquivo de tags carregado com sucesso.')
            return self.mapa_tags
        except FileNotFoundError:
            raise FileNotFoundError(f"Arquivo de tags '{arquivo_tags}' não encontrado.")
        except json.JSONDecodeError as e:
            raise Exception(f"Erro ao carregar o arquivo de tags '{arquivo_tags}': {e}")
        except Exception as e:
            raise Exception(f"Erro ao carregar o arquivo de tags '{arquivo_tags}': {e}")

    def substituir_tags(self):
        if self.mapa_tags is None:
            raise ValueError('Mapa de tags não foi carregado. Carregue um arquivo de tags com carregar_tags().')

        for tag, valor in self.mapa_tags.items():
            if valor is not None:
                if tag.startswith('|IMAGEM'):
                    self.substituir_grafico(tag, valor)
                elif tag.startswith('|TABELA'):
                    self.substituir_tabela(tag, valor)
                else:
                    self.substituir_texto(tag, valor)

    def substituir_texto(self, tag, valor):
        for paragraph in self.doc.paragraphs:
            if tag in paragraph.text:
                for run in paragraph.runs:
                    if tag in run.text:
                        run.text = run.text.replace(tag, valor)

    def substituir_grafico(self, tag, valor):
        for i, paragraph in enumerate(self.doc.paragraphs):
            if tag in paragraph.text:
                paragraph.text = paragraph.text.replace(tag, '')
                run = paragraph.add_run()
                run.add_picture(valor, width=Inches(6.0))

    def substituir_tabela(self, tag, valor):
        try:
            dataframe = pd.read_csv(valor, sep=',', encoding='utf-8')
        except Exception as e:
            raise f"Erro ao carregar o arquivo de dados: {e},  *dataframe = pd.read_csv(valor, sep=',', encoding='utf-8'))"


        for paragraph in self.doc.paragraphs:
            if tag in paragraph.text:
                # Substitui a tag pela string vazia
                paragraph.text = paragraph.text.replace(tag, '')

                # Cria a tabela
                table = self.doc.add_table(rows=dataframe.shape[0] + 1, cols=dataframe.shape[1])
                table.alignment = 1  # Centraliza a tabela no documento (opcional)

                # Preenche o cabeçalho
                for j, col in enumerate(dataframe.columns):
                    table.cell(0, j).text = col

                # Preenche as células da tabela
                for i, row in dataframe.iterrows():
                    for j, col in enumerate(dataframe.columns):
                        table.cell(i + 1, j).text = str(row[col])

                # Move a tabela para o local correto
                table_element = table._element
                paragraph._element.addnext(table_element)
                break  # Sai do loop após inserir a tabela

    def gerar_relatorio(self, arquivo_saida):
        self.substituir_tags()
        self.doc.save(arquivo_saida)
